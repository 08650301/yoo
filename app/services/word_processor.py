import os
import re
import mammoth
from docx import Document
from io import BytesIO
from docxcompose.composer import Composer
from app import db
from app.models import Project, Template, FixedFormData, DYNAMIC_TABLE_MODELS, Section, SheetDefinition, WordTemplateChapter

# --- Private Helper Functions ---

def _replace_text_in_paragraph(paragraph, key, value):
    """在段落中替换文本占位符"""
    # 简单的替换，对于复杂的格式可能会有问题
    if key in paragraph.text:
        inline = paragraph.runs
        # 替换段落中的文本，同时尽量保留格式
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, str(value if value is not None else ''))
                inline[i].text = text

def _replace_placeholders_in_doc(doc, placeholders):
    """替换文档中的所有文本占位符"""
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            # 确保不是表格占位符
            if not key.startswith('{{table_'):
                 _replace_text_in_paragraph(p, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in placeholders.items():
                        if not key.startswith('{{table_'):
                            _replace_text_in_paragraph(p, key, value)

def _replace_table_placeholder(doc, placeholder_text, table_data, column_config):
    """查找并替换表格占位符"""
    for p in doc.paragraphs:
        if placeholder_text in p.text:
            # 清空占位符所在的段落
            p.clear()

            if not table_data:
                # 如果没有数据，可以选择不添加表格或添加一个空表头
                p.text = "(此部分无数据)" # 或者直接返回
                return True

            headers = [col['label'] for col in column_config]
            # 在占位符段落的位置插入新表格
            table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')

            # 填充表头
            hdr_cells = table.rows[0].cells
            for i, header_name in enumerate(headers):
                hdr_cells[i].text = header_name

            # 填充数据行
            for item in table_data:
                row_cells = table.add_row().cells
                for i, col_config in enumerate(column_config):
                    cell_value = item.get(col_config['name'], '')
                    row_cells[i].text = str(cell_value if cell_value is not None else '')
            return True
    return False

# --- Public Service Functions ---

def generate_preview_html(project):
    """为指定项目生成所有关联章节的Word模板的HTML预览拼接"""
    template = Template.query.filter_by(name=project.procurement_method, is_latest=True).first()
    if not template:
        raise ValueError("未找到已发布的模板")

    # 新的查询逻辑：从 Template 开始，join Section 和 SheetDefinition
    sheets_with_chapters = db.session.query(SheetDefinition).join(Section).filter(
        Section.template_id == template.id,
        SheetDefinition.word_template_chapter_id.isnot(None)
    ).order_by(Section.display_order, SheetDefinition.display_order).all()

    if not sheets_with_chapters:
        return "<p class='text-danger'>此模板下没有任何Sheet关联了章节文档，无法生成预览。</p>"

    full_html = ""
    for sheet in sheets_with_chapters:
        chapter_path = sheet.word_template_chapter.filepath
        if not os.path.exists(chapter_path):
            full_html += f"<p class='text-danger'>错误: 章节 '{sheet.name}' 的模板文件不存在: {chapter_path}</p><hr>"
            continue

        with open(chapter_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

            def wrap_placeholder(match):
                placeholder = match.group(1)
                # 保持与 live_preview.js 的兼容性，只使用字段名作为 key
                return f'<span data-placeholder-for="{placeholder}">{match.group(0)}</span>'

            html = re.sub(r'\{\{([\w_]+)\}\}', wrap_placeholder, html)
            full_html += html + "<hr>"

    return full_html


def generate_word_document(project, template, template_config):
    """为指定项目生成最终的Word文档（通过合并章节）"""

    # 1. 获取所有固定表单数据，构建一个大的占位符字典
    placeholders = {}
    fixed_data = FixedFormData.query.filter_by(project_id=project.id).all()
    for item in fixed_data:
        placeholders[f"{{{{{item.field_name}}}}}"] = item.field_value

    # 2. 按正确顺序找到所有关联了章节文档的Sheet
    sheets_with_chapters = db.session.query(SheetDefinition).join(Section).filter(
        Section.template_id == template.id,
        SheetDefinition.word_template_chapter_id.isnot(None)
    ).order_by(Section.display_order, SheetDefinition.display_order).all()

    if not sheets_with_chapters:
        # 如果没有任何章节关联，创建一个提示错误的文档
        doc = Document()
        doc.add_paragraph("错误：此模板下没有任何Sheet关联了章节文档，无法生成文档。")
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    # 3. 创建主文档 (基于第一个章节模板)
    first_sheet = sheets_with_chapters[0]
    first_chapter_path = first_sheet.word_template_chapter.filepath
    if not os.path.exists(first_chapter_path):
        raise FileNotFoundError(f"起始章节模板文件不存在: {first_chapter_path}")

    master_doc = Document(first_chapter_path)
    composer = Composer(master_doc)

    # 4. 遍历所有章节，填充并合并
    # 第一个章节已经作为主文档加载，所以我们从它开始处理
    for i, sheet in enumerate(sheets_with_chapters):
        chapter_path = sheet.word_template_chapter.filepath
        if not os.path.exists(chapter_path):
            # 在文档中插入一个警告，而不是让整个过程失败
            master_doc.add_paragraph(f"警告：章节 '{sheet.name}' 的模板文件未找到，已跳过。")
            continue

        # 如果是第一个文档，我们直接在 composer 的主文档上操作
        # 如果是后续文档，则加载它并追加
        if i == 0:
            doc_to_process = master_doc
        else:
            doc_to_process = Document(chapter_path)

        # 填充文本占位符
        _replace_placeholders_in_doc(doc_to_process, placeholders)

        # 填充表格占位符
        if sheet.sheet_type == "dynamic_table":
            model_identifier = sheet.model_identifier
            table_placeholder = f"{{{{table_{model_identifier}}}}}"
            Model = DYNAMIC_TABLE_MODELS.get(model_identifier)
            if Model:
                table_data = Model.query.filter_by(project_id=project.id).all()
                table_data_dicts = [dict((col, getattr(d, col)) for col in d.__table__.columns.keys()) for d in table_data]

                # 获取该动态表格的列定义
                column_config = next((form['columns'] for sec in template_config.get("sections", {}).values() for form_name, form in sec.get("forms", {}).items() if form_name == sheet.name), [])
                _replace_table_placeholder(doc_to_process, table_placeholder, table_data_dicts, column_config)

        # 如果不是第一个文档，则将其追加到主文档
        if i > 0:
            composer.append(doc_to_process)

    # 5. 保存最终文档到内存流
    file_stream = BytesIO()
    composer.save(file_stream)
    file_stream.seek(0)
    return file_stream