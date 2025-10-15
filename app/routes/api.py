import pandas as pd
import mammoth
import os
import re
from docx import Document
from io import BytesIO
from datetime import timezone, timedelta
from flask import Blueprint, jsonify, request, send_file
from app import db
from app.models import (
    Project, Template, Section, SheetDefinition, FieldDefinition, ConditionalRule,
    FixedFormData, DYNAMIC_TABLE_MODELS
)

api_bp = Blueprint('api', __name__, url_prefix='/api')

# ... (existing helper functions like get_config_from_db, find_sheet_config_from_db) ...
# ... (existing API routes like get_forms_config_api, projects CRUD, etc.) ...

def get_config_from_db(procurement_method):
    """
    根据采购方式（模板名称）从数据库动态生成前端所需的配置JSON。
    """
    template = Template.query.filter_by(name=procurement_method, status='published', is_latest=True).first()
    if not template:
        return None

    config = {"sections": {}}
    sections_query = Section.query.filter_by(template_id=template.id).order_by(Section.display_order).all()

    for section in sections_query:
        section_config = {"order": [], "forms": {}}
        sheets_query = SheetDefinition.query.filter_by(section_id=section.id).order_by(
            SheetDefinition.display_order).all()

        for sheet in sheets_query:
            section_config["order"].append(sheet.name)
            sheet_config = {
                "type": sheet.sheet_type,
                "model_identifier": sheet.model_identifier
            }
            fields_query = FieldDefinition.query.filter_by(sheet_id=sheet.id).order_by(
                FieldDefinition.display_order).all()

            # 为前端传递完整的字段信息，包括校验规则和选项
            fields_list = []
            for f in fields_query:
                fields_list.append({
                    "name": f.name,
                    "label": f.label,
                    "type": f.field_type,
                    "default_value": f.default_value,
                    "options": f.options,
                    "validation_rules": [{"rule_type": r.rule_type, "rule_value": r.rule_value} for r in
                                         f.validation_rules]
                })

            if sheet.sheet_type == 'fixed_form':
                sheet_config['fields'] = fields_list
                # 获取并添加联动规则
                rules = ConditionalRule.query.filter_by(sheet_id=sheet.id).order_by(ConditionalRule.id).all()
                sheet_config['conditional_rules'] = [
                    {"id": r.id, "name": r.name, "definition": r.definition} for r in rules
                ]
            else:
                sheet_config['columns'] = fields_list

            section_config["forms"][sheet.name] = sheet_config

        config["sections"][section.name] = section_config

    return config


def find_sheet_config_from_db(procurement_method, sheet_name):
    """
    根据采购方式和Sheet名称，查找单个Sheet的配置。
    """
    template = Template.query.filter_by(name=procurement_method, is_latest=True).first()
    if not template:
        return None

    sheet_def = SheetDefinition.query.join(Section).filter(
        Section.template_id == template.id,
        SheetDefinition.name == sheet_name
    ).first()

    if not sheet_def:
        return None

    config = {"type": sheet_def.sheet_type, "model_identifier": sheet_def.model_identifier}
    fields_query = FieldDefinition.query.filter_by(sheet_id=sheet_def.id).order_by(FieldDefinition.display_order).all()
    fields_list = [{"name": f.name, "label": f.label, "type": f.field_type, "default_value": f.default_value} for f in fields_query]

    if sheet_def.sheet_type == 'fixed_form':
        config['fields'] = fields_list
    else:
        config['columns'] = fields_list

    return config

# ==============================================================================
# 前端 API 路由
# ==============================================================================

@api_bp.route('/forms-config/<string:method>')
def get_forms_config_api(method):
    config = get_config_from_db(method)
    if config:
        return jsonify(config)
    return jsonify({"error": "未知的或未发布的采购方式"}), 404


@api_bp.route('/published-templates')
def get_published_templates():
    templates = Template.query.filter_by(status='published', is_latest=True).order_by(Template.display_order).all()
    return jsonify([t.name for t in templates])


@api_bp.route('/projects', methods=['GET'])
def get_projects():
    try:
        query = Project.query
        name_query = request.args.get('name', '').strip()
        number_query = request.args.get('number', '').strip()
        method_query = request.args.get('method', '').strip()

        if name_query:
            query = query.filter(Project.name.like(f"%{name_query}%"))
        if number_query:
            query = query.filter(Project.number.like(f"%{number_query}%"))
        if method_query:
            query = query.filter(Project.procurement_method == method_query)

        projects = query.order_by(Project.created_at.desc()).all()
        china_tz = timezone(timedelta(hours=8))

        return jsonify([{
            "id": p.id,
            "name": p.name,
            "number": p.number,
            "procurement_method": p.procurement_method,
            "created_at": p.created_at.replace(tzinfo=timezone.utc).astimezone(china_tz).strftime('%Y-%m-%d %H:%M')
        } for p in projects])
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@api_bp.route('/projects', methods=['POST'])
def create_project():
    data = request.json
    if not data or not data.get('name', '').strip():
        return jsonify({"error": "项目名称不能为空"}), 400
    if not data.get('number', '').strip():
        return jsonify({"error": "项目编号不能为空"}), 400
    if not data.get('procurement_method'):
        return jsonify({"error": "必须选择采购方式"}), 400

    new_project = Project(name=data['name'], number=data['number'], procurement_method=data['procurement_method'])
    db.session.add(new_project)
    db.session.commit()
    return jsonify({"id": new_project.id, "name": new_project.name, "message": "项目创建成功"}), 201


@api_bp.route('/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    try:
        project = Project.query.get_or_404(project_id)
        FixedFormData.query.filter_by(project_id=project_id).delete()
        for model in DYNAMIC_TABLE_MODELS.values():
            model.query.filter_by(project_id=project_id).delete()
        db.session.delete(project)
        db.session.commit()
        return jsonify({"message": "项目已成功删除"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@api_bp.route('/projects/<int:project_id>', methods=['PUT'])
def update_project(project_id):
    try:
        project = Project.query.get_or_404(project_id)
        data = request.json
        if not data or not data.get('name', '').strip():
            return jsonify({"error": "项目名称不能为空"}), 400
        if not data.get('number', '').strip():
            return jsonify({"error": "项目编号不能为空"}), 400
        if not data.get('procurement_method'):
            return jsonify({"error": "必须选择采购方式"}), 400

        project.name = data['name']
        project.number = data['number']
        project.procurement_method = data['procurement_method']
        db.session.commit()
        return jsonify({"message": "项目信息更新成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@api_bp.route('/projects/<int:project_id>/sheets/<string:sheet_name>', methods=['GET'])
def get_sheet_data(project_id, sheet_name):
    project = Project.query.get_or_404(project_id)
    config = find_sheet_config_from_db(project.procurement_method, sheet_name)
    if not config:
        return jsonify({"error": "Sheet名称不存在"}), 404

    if config['type'] == 'fixed_form':
        return jsonify({entry.field_name: entry.field_value for entry in
                        FixedFormData.query.filter_by(project_id=project_id, sheet_name=sheet_name).all()})
    elif config['type'] == 'dynamic_table':
        model_identifier = config.get('model_identifier')
        Model = DYNAMIC_TABLE_MODELS.get(model_identifier)
        if not Model:
            return jsonify({"error": f"未找到标识符为 {model_identifier} 的模型"}), 404

        items = Model.query.filter_by(project_id=project_id).order_by(Model.id).all()
        return jsonify(
            [{c.name: getattr(item, c.name) for c in item.__table__.columns if c.name not in ['id', 'project_id']} for
             item in items]
        )


@api_bp.route('/projects/<int:project_id>/sheets/<string:sheet_name>', methods=['POST'])
def save_sheet_data(project_id, sheet_name):
    project = Project.query.get_or_404(project_id)
    config = find_sheet_config_from_db(project.procurement_method, sheet_name)
    if not config:
        return jsonify({"error": "Sheet名称不存在"}), 404

    data = request.json
    try:
        if config['type'] == 'fixed_form':
            FixedFormData.query.filter_by(project_id=project_id, sheet_name=sheet_name).delete()
            for field_name, field_value in data.items():
                if any(f['name'] == field_name for f in config.get('fields', [])):
                    db.session.add(FixedFormData(project_id=project_id, sheet_name=sheet_name, field_name=field_name,
                                                 field_value=str(field_value)))
        elif config['type'] == 'dynamic_table':
            model_identifier = config.get('model_identifier')
            Model = DYNAMIC_TABLE_MODELS.get(model_identifier)
            if not Model:
                return jsonify({"error": f"未找到标识符为 {model_identifier} 的模型"}), 404

            Model.query.filter_by(project_id=project_id).delete()
            valid_columns = {c.name: c.type for c in Model.__table__.columns if c.name not in ['id', 'project_id']}
            for row_data in data:
                processed_row = {
                    key: (None if isinstance(valid_columns.get(key), (db.Float, db.Integer)) and value == '' else value)
                    for key, value in row_data.items() if key in valid_columns
                }
                data_without_seq = processed_row.copy()
                data_without_seq.pop('seq_num', None)
                if not all(val is None or str(val).strip() == '' for val in data_without_seq.values()):
                    db.session.add(Model(project_id=project_id, **processed_row))

        db.session.commit()
        return jsonify({"message": f"'{sheet_name}' 保存成功"})
    except Exception as e:
        db.session.rollback()
        # print(f"保存失败! Sheet: {sheet_name}, 错误: {e}")
        return jsonify({"error": str(e)}), 500


@api_bp.route('/projects/<int:project_id>/export/<string:section_name>')
def export_project_excel(project_id, section_name):
    # ... (existing code for excel export)
    pass

@api_bp.route('/projects/<int:project_id>/preview', methods=['GET'])
def get_word_preview(project_id):
    # ... (existing code for preview)
    pass

# ==============================================================================
# Word 导出相关辅助函数和路由
# ==============================================================================

def replace_text_in_paragraph(paragraph, key, value):
    """在段落中替换文本占位符"""
    # Simple replacement
    if key in paragraph.text:
        inline = paragraph.runs
        # Replace strings and retain formatting
        for i in range(len(inline)):
            if key in inline[i].text:
                text = inline[i].text.replace(key, str(value))
                inline[i].text = text

def replace_placeholders_in_doc(doc, placeholders):
    """替换文档中的所有文本和表格占位符"""
    # 替换段落中的文本
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            # 跳过表格占位符
            if not key.startswith('{{table_'):
                 replace_text_in_paragraph(p, key, value)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in placeholders.items():
                        if not key.startswith('{{table_'):
                            replace_text_in_paragraph(p, key, value)

def replace_table_placeholder(doc, placeholder_text, table_data, column_config):
    """查找并替换表格占位符"""
    for p in doc.paragraphs:
        if placeholder_text in p.text:
            p.text = "" # 清空段落
            # 在该段落处插入表格
            headers = [col['label'] for col in column_config]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header_name in enumerate(headers):
                hdr_cells[i].text = header_name

            for item in table_data:
                row_cells = table.add_row().cells
                for i, col_config in enumerate(column_config):
                    cell_value = item.get(col_config['name'], '')
                    row_cells[i].text = str(cell_value if cell_value is not None else '')
            return True # 表示已找到并替换
    return False


@api_bp.route('/projects/<int:project_id>/export_word', methods=['GET'])
def export_word_document(project_id):
    project = Project.query.get_or_404(project_id)
    template = Template.query.filter_by(name=project.procurement_method, is_latest=True).first()

    if not template or not template.word_template_path or not os.path.exists(template.word_template_path):
        return jsonify({"error": "未找到或未关联有效的Word模板文件"}), 404

    try:
        doc = Document(template.word_template_path)

        # 1. 准备所有占位符的数据
        placeholders = {}

        # a. 从 FixedFormData 收集
        fixed_data = FixedFormData.query.filter_by(project_id=project_id).all()
        for item in fixed_data:
            placeholders[f"{{{{{item.field_name}}}}}"] = item.field_value

        # b. 替换文档中的文本占位符
        replace_placeholders_in_doc(doc, placeholders)

        # c. 查找并替换所有动态表格占位符
        template_config = get_config_from_db(project.procurement_method)
        if template_config:
            for section_config in template_config.get("sections", {}).values():
                for sheet_name, form_config in section_config.get("forms", {}).items():
                    if form_config.get("type") == "dynamic_table":
                        model_identifier = form_config.get("model_identifier")
                        table_placeholder = f"{{{{table_{model_identifier}}}}}"
                        Model = DYNAMIC_TABLE_MODELS.get(model_identifier)
                        if Model:
                            table_data = Model.query.filter_by(project_id=project_id).all()
                            table_data_dicts = [dict((col, getattr(d, col)) for col in d.__table__.columns.keys()) for d in table_data]
                            replace_table_placeholder(doc, table_placeholder, table_data_dicts, form_config.get('columns', []))

        # 3. 保存到内存并发送
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        safe_project_name = "".join([c for c in project.name if c.isalnum() or c in (' ', '-')]).rstrip()
        filename = f"{safe_project_name}_导出.docx"

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({"error": f"生成Word文档时出错: {e}"}), 500
